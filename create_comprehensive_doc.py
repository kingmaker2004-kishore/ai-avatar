from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_comprehensive_word_document():
    """Create one comprehensive Word document with all documentation"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Title page
    title = doc.add_heading('AI Avatar', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Complete Creator Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(18)
    subtitle_run.bold = True
    
    doc.add_paragraph()
    date_para = doc.add_paragraph('April 2026')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    version_para = doc.add_paragraph('Version 1.0')
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Table of contents
    doc.add_page_break()
    toc_heading = doc.add_heading('Table of Contents', level=1)
    
    toc_items = [
        'Project Overview',
        'Developer Setup',
        'System Architecture',
        'Backend API Reference',
        'Frontend Components',
        'User Guide',
        'Production Deployment',
        'Troubleshooting',
        'Contributing Guidelines'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # Read and add complete documentation
    doc.add_page_break()
    
    with open('docs/COMPLETE_DOCUMENTATION.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Parse and add content
    lines = content.split('\n')
    
    for i, line in enumerate(lines):
        # Skip the first table of contents section
        if i < 100 and ('# Table of Contents' in line or 'docs/' in line):
            continue
        
        # Skip empty lines at beginning
        if i < 50 and not line.strip():
            continue
        
        # Process markdown formatting
        if line.startswith('# '):
            # Main heading
            title = line[2:].strip()
            if '##' not in line:
                doc.add_page_break()
                doc.add_heading(title, level=1)
        
        elif line.startswith('## '):
            # Subheading
            doc.add_heading(line[3:].strip(), level=2)
        
        elif line.startswith('### '):
            # Sub-subheading
            doc.add_heading(line[4:].strip(), level=3)
        
        elif line.startswith('#### '):
            # Minor heading
            doc.add_heading(line[5:].strip(), level=4)
        
        elif line.strip().startswith('```'):
            # Code block - collect until closing ```
            code_lines = []
            j = i + 1
            while j < len(lines) and not lines[j].strip().startswith('```'):
                code_lines.append(lines[j])
                j += 1
            
            if code_lines:
                code_text = '\n'.join(code_lines).strip()
                code_para = doc.add_paragraph(code_text, style='List Bullet')
                for run in code_para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
        
        elif line.startswith('| '):
            # Table row - for now just add as text
            doc.add_paragraph(line)
        
        elif line.startswith('- '):
            # Bullet point
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        
        elif line.startswith('**') and line.endswith('**'):
            # Bold text
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2].strip())
            run.bold = True
        
        elif line.strip():
            # Regular paragraph
            doc.add_paragraph(line.strip())
        
        elif i > 100:  # After main content starts, add spacing
            doc.add_paragraph()
    
    # Save the comprehensive document
    output_path = os.path.join('doc', 'AI_Avatar_Complete_Documentation.docx')
    
    # Create doc folder if it doesn't exist
    os.makedirs('doc', exist_ok=True)
    
    doc.save(output_path)
    print(f'✅ Comprehensive documentation created: {output_path}')
    print(f'   File size: {os.path.getsize(output_path) / 1024:.1f} KB')
    
    return output_path

if __name__ == '__main__':
    print('📄 Creating comprehensive Word document...\n')
    create_comprehensive_word_document()
    print('\n✨ Complete! One comprehensive Word document has been created.\n')
    print('File location: doc/AI_Avatar_Complete_Documentation.docx')
    print('This single document contains all sections:')
    print('  • Project Overview')
    print('  • Developer Setup')
    print('  • System Architecture')
    print('  • Backend API Reference')
    print('  • Frontend Components')
    print('  • User Guide')
    print('  • Production Deployment')
    print('  • Troubleshooting')
    print('  • Contributing Guidelines')
