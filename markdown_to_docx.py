from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from pathlib import Path

def add_heading_from_markdown(doc, text, level=1):
    """Add a heading based on markdown level"""
    style = f'Heading {level}'
    doc.add_heading(text, level=level)

def read_markdown_file(filepath):
    """Read markdown file and return content"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def parse_markdown_to_docx(markdown_content):
    """Parse markdown content and return formatted paragraphs"""
    lines = markdown_content.split('\n')
    paragraphs = []
    current_code_block = []
    in_code_block = False
    
    for line in lines:
        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                paragraphs.append({
                    'type': 'code',
                    'content': '\n'.join(current_code_block)
                })
                current_code_block = []
                in_code_block = False
            else:
                in_code_block = True
                current_code_block = []
            continue
        
        if in_code_block:
            current_code_block.append(line)
            continue
        
        # Headings
        if line.startswith('# '):
            paragraphs.append({'type': 'heading1', 'content': line[2:].strip()})
        elif line.startswith('## '):
            paragraphs.append({'type': 'heading2', 'content': line[3:].strip()})
        elif line.startswith('### '):
            paragraphs.append({'type': 'heading3', 'content': line[4:].strip()})
        elif line.startswith('#### '):
            paragraphs.append({'type': 'heading4', 'content': line[5:].strip()})
        # Bold text
        elif line.startswith('**') and line.endswith('**'):
            paragraphs.append({'type': 'bold', 'content': line[2:-2]})
        # Regular paragraph
        elif line.strip():
            paragraphs.append({'type': 'paragraph', 'content': line.strip()})
        # Blank line
        elif not line.strip():
            if paragraphs and paragraphs[-1]['type'] != 'blank':
                paragraphs.append({'type': 'blank', 'content': ''})
    
    return paragraphs

def add_content_to_doc(doc, paragraphs):
    """Add parsed content to document"""
    for para in paragraphs:
        ptype = para['type']
        content = para['content']
        
        if ptype == 'heading1':
            doc.add_heading(content, level=1)
        elif ptype == 'heading2':
            doc.add_heading(content, level=2)
        elif ptype == 'heading3':
            doc.add_heading(content, level=3)
        elif ptype == 'heading4':
            doc.add_heading(content, level=4)
        elif ptype == 'bold':
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
        elif ptype == 'code':
            p = doc.add_paragraph(content, style='List Bullet')
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        elif ptype == 'paragraph':
            doc.add_paragraph(content)
        elif ptype == 'blank':
            doc.add_paragraph()

def create_master_document():
    """Create a master Word document with all documentation"""
    doc = Document()
    
    # Title page
    title = doc.add_heading('AI Avatar', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Complete Creator Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(18)
    
    doc.add_paragraph('April 2026')
    
    # Table of Contents
    doc.add_page_break()
    doc.add_heading('Table of Contents', level=1)
    
    toc_items = [
        ('README - Project Overview', 'docs/README.md'),
        ('SETUP - Developer Environment', 'docs/SETUP.md'),
        ('ARCHITECTURE - System Design', 'docs/ARCHITECTURE.md'),
        ('API - Backend Reference', 'docs/API.md'),
        ('COMPONENTS - Frontend Documentation', 'docs/COMPONENTS.md'),
        ('USER_GUIDE - End-User Guide', 'docs/USER_GUIDE.md'),
        ('DEPLOYMENT - Production Setup', 'docs/DEPLOYMENT.md'),
        ('TROUBLESHOOTING - Common Issues', 'docs/TROUBLESHOOTING.md'),
        ('CONTRIBUTING - Developer Guidelines', 'docs/CONTRIBUTING.md'),
    ]
    
    for title, path in toc_items:
        doc.add_paragraph(title, style='List Bullet')
    
    # Add each document
    docs_to_include = [
        'docs/README.md',
        'docs/SETUP.md',
        'docs/ARCHITECTURE.md',
        'docs/API.md',
        'docs/COMPONENTS.md',
        'docs/USER_GUIDE.md',
        'docs/DEPLOYMENT.md',
        'docs/TROUBLESHOOTING.md',
        'docs/CONTRIBUTING.md',
    ]
    
    for doc_path in docs_to_include:
        if os.path.exists(doc_path):
            doc.add_page_break()
            content = read_markdown_file(doc_path)
            paragraphs = parse_markdown_to_docx(content)
            add_content_to_doc(doc, paragraphs)
    
    # Save document
    output_path = 'AI_Avatar_Documentation.docx'
    doc.save(output_path)
    print(f'✅ Master documentation created: {output_path}')
    return output_path

def create_individual_documents():
    """Create individual Word documents for each guide"""
    docs_to_create = {
        'AI_Avatar_README.docx': 'docs/README.md',
        'AI_Avatar_SETUP.docx': 'docs/SETUP.md',
        'AI_Avatar_ARCHITECTURE.docx': 'docs/ARCHITECTURE.md',
        'AI_Avatar_API.docx': 'docs/API.md',
        'AI_Avatar_COMPONENTS.docx': 'docs/COMPONENTS.md',
        'AI_Avatar_USER_GUIDE.docx': 'docs/USER_GUIDE.md',
        'AI_Avatar_DEPLOYMENT.docx': 'docs/DEPLOYMENT.md',
        'AI_Avatar_TROUBLESHOOTING.docx': 'docs/TROUBLESHOOTING.md',
        'AI_Avatar_CONTRIBUTING.docx': 'docs/CONTRIBUTING.md',
    }
    
    for output_filename, input_path in docs_to_create.items():
        if os.path.exists(input_path):
            doc = Document()
            content = read_markdown_file(input_path)
            paragraphs = parse_markdown_to_docx(content)
            add_content_to_doc(doc, paragraphs)
            doc.save(output_filename)
            print(f'✅ Created: {output_filename}')

if __name__ == '__main__':
    print('📄 Converting AI Avatar documentation to Word format...\n')
    
    # Create master document
    create_master_document()
    
    print('\n📄 Creating individual documents...\n')
    
    # Create individual documents
    create_individual_documents()
    
    print('\n✨ All Word documents created successfully!')
    print('\nFiles created:')
    print('  • AI_Avatar_Documentation.docx (Master - all guides combined)')
    print('  • AI_Avatar_README.docx')
    print('  • AI_Avatar_SETUP.docx')
    print('  • AI_Avatar_ARCHITECTURE.docx')
    print('  • AI_Avatar_API.docx')
    print('  • AI_Avatar_COMPONENTS.docx')
    print('  • AI_Avatar_USER_GUIDE.docx')
    print('  • AI_Avatar_DEPLOYMENT.docx')
    print('  • AI_Avatar_TROUBLESHOOTING.docx')
    print('  • AI_Avatar_CONTRIBUTING.docx')
